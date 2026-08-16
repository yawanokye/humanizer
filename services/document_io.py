from __future__ import annotations

import io
import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Pt
from pypdf import PdfReader

SUPPORTED_EXTENSIONS = {".txt", ".md", ".docx", ".pdf"}


def extract_text(filename: str, content: bytes) -> str:
    suffix = Path(filename or "").suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError("Supported files are TXT, MD, DOCX and text-based PDF.")
    if suffix in {".txt", ".md"}:
        for encoding in ("utf-8-sig", "utf-8", "cp1252"):
            try:
                return content.decode(encoding)
            except UnicodeDecodeError:
                continue
        raise ValueError("The text file encoding could not be read.")
    if suffix == ".docx":
        document = Document(io.BytesIO(content))
        blocks: list[str] = []
        for paragraph in document.paragraphs:
            blocks.append(paragraph.text)
        for table in document.tables:
            for row in table.rows:
                blocks.append(" | ".join(cell.text.strip() for cell in row.cells))
        return "\n\n".join(block for block in blocks if block.strip())
    reader = PdfReader(io.BytesIO(content))
    pages = [(page.extract_text() or "").strip() for page in reader.pages]
    text = "\n\n".join(page for page in pages if page)
    if not text.strip():
        raise ValueError("No selectable text was found in the PDF. OCR is not included in this standalone build.")
    return text


def _heading_level(line: str) -> int | None:
    clean = line.strip()
    match = re.match(r"^(#{1,6})\s+", clean)
    if match:
        return min(4, len(match.group(1)))
    if re.match(r"^CHAPTER\s+(?:\d+|[A-Z]+)", clean, re.I):
        return 1
    number = re.match(r"^(\d+(?:\.\d+){1,3})\s+", clean)
    if number:
        return min(4, number.group(1).count(".") + 1)
    return None


def build_docx(text: str, title: str = "Scholarly Humanized Text") -> bytes:
    document = Document()
    document.core_properties.title = title
    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(11)
    if title:
        document.add_heading(title, level=0)
    for block in re.split(r"\n\s*\n", str(text or "")):
        clean = block.strip()
        if not clean:
            continue
        level = _heading_level(clean)
        if level is not None and "\n" not in clean:
            document.add_heading(re.sub(r"^#{1,6}\s+", "", clean), level=level)
        else:
            document.add_paragraph(clean)
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def build_annotated_docx(text: str, segments: list[dict[str, Any]], title: str = "AI Signal Diagnostic") -> bytes:
    document = Document()
    document.core_properties.title = title
    document.add_heading(title, level=0)
    document.add_paragraph(
        "AI signal key: red = high signal, dark yellow = moderate, yellow = low, uncoloured = minimal signal or protected. "
        "The diagnostic estimates AI-like writing patterns and does not prove authorship."
    )
    by_paragraph: dict[int, list[dict[str, Any]]] = {}
    for segment in segments:
        by_paragraph.setdefault(int(segment["paragraph_index"]), []).append(segment)
    for paragraph_index in sorted(by_paragraph):
        paragraph = document.add_paragraph()
        for segment in by_paragraph[paragraph_index]:
            run = paragraph.add_run(segment["text"])
            band = segment["band"]
            if band == "high":
                run.font.highlight_color = WD_COLOR_INDEX.RED
            elif band == "moderate":
                run.font.highlight_color = WD_COLOR_INDEX.DARK_YELLOW
            elif band == "low":
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            run.font.size = Pt(11)
            paragraph.add_run(" ")
    document.add_page_break()
    document.add_heading("Sentence-level AI-style evidence", level=1)
    for segment in segments:
        if segment["band"] in {"high", "moderate", "low"}:
            paragraph = document.add_paragraph(style="List Number")
            paragraph.add_run(f'{segment["risk"]}% concern: ').bold = True
            paragraph.add_run("; ".join(segment.get("reasons") or []))
            quote = document.add_paragraph(segment["text"].strip())
            quote.style = document.styles["Intense Quote"]
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()
